"""Document processor for handling multiple file types."""

from pathlib import Path
from typing import Optional
from abc import ABC, abstractmethod
from dataclasses import dataclass
import numpy as np
from dotenv import load_dotenv
import os
load_dotenv()
import subprocess

from fs_rag.core import get_logger , get_config
config = get_config()

logger = get_logger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of document text."""
    content: str
    source_file: Path
    chunk_index: int
    metadata: dict


def format_path(path):
    path = str(path)
    result = ""

    prefix = config.filepath_prefix_to_remove 
    path = path.replace(f"{prefix}", "")

    for ch in path:
        if ch == "/":
            result += " "
        else:
            result += ch
    result += "  "
    return result

class DocumentProcessor(ABC):
    """Base class for document processors."""

    @abstractmethod
    def can_process(self, file_path: Path) -> bool:
        """Check if this processor can handle the file."""
        pass

    @abstractmethod
    def extract_text(self, file_path: Path) -> str:
        """Extract text from the document."""
        pass

    def chunk_text(self, file_path: str,text: str, chunk_size: int = 512, chunk_overlap: int = 50) -> list[str]:
        """Split text into overlapping chunks."""
        chunks = []
        if len(text) <= chunk_size:
            return [text]

        formated_file_path = ''
        if config.enable_filepath_injection:     
            formated_file_path = format_path(file_path)

        for i in range(0, len(text), chunk_size - chunk_overlap):
            chunk = formated_file_path + text[i : i + chunk_size]
            if chunk.strip():
                chunks.append(chunk)        
        return chunks


class TextProcessor(DocumentProcessor):
    """Processor for plain text files."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".txt", ".md", ".log"}

    def extract_text(self, file_path: Path) -> str:
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            logger.warning(f"Failed to decode {file_path} as UTF-8, trying latin-1")
            return file_path.read_text(encoding="latin-1")


class PDFProcessor(DocumentProcessor):
    """Processor for PDF files with OCR fallback for scanned PDFs."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"

    def _extract_with_pypdf2(self, file_path: Path) -> tuple[str, bool]:
        """
        Try to extract text using PyPDF2.
        Returns: (extracted_text, success_flag)
        """
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            logger.debug("PyPDF2 not installed")
            return "", False

        try:
            text = []
            with open(file_path, "rb") as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted and extracted.strip():
                        text.append(extracted)
            
            combined_text = "\n".join(text)
            
            # Check if extraction was successful (at least 10 chars)
            if len(combined_text.strip()) > 10:
                logger.debug(f"PyPDF2 extraction successful: {len(combined_text)} chars")
                return combined_text, True
            else:
                logger.debug(f"PyPDF2 extraction returned minimal text, likely scanned PDF")
                return "", False
                
        except Exception as e:
            logger.debug(f"PyPDF2 extraction failed for {file_path}: {e}")
            return "", False

    def _extract_with_ocr(self, file_path: Path) -> str:
        """
        Extract text from PDF using OCR (fallback for scanned PDFs).
        Uses pdf2image + EasyOCR for performance.
        """
        try:
            from pdf2image import convert_from_path
        except ImportError:
            logger.warning("pdf2image not installed. Cannot perform OCR.")
            return ""

        try:
            import easyocr
        except ImportError:
            logger.warning("EasyOCR not installed. Cannot perform OCR.")
            return ""

        try:
            logger.info(f"Starting OCR extraction for {file_path.name}")
            
            # Convert PDF pages to images (limit to first 50 pages for performance)
            images = convert_from_path(str(file_path), first_page=1, last_page=50)
            
            if not images:
                logger.warning(f"No images extracted from {file_path}")
                return ""
            
            logger.debug(f"Converted {len(images)} PDF pages to images")
            
            def str_to_bool(value: str) -> bool:
                return str(value).strip().lower() in ("1", "true", "yes", "on")

            use_gpu = str_to_bool(os.getenv("OCR_USE_GPU", "false"))

            # Initialize OCR reader (lazy load)
            reader = easyocr.Reader(['pt', 'en'], gpu=use_gpu)
            
            all_text = []
            for page_num, image in enumerate(images, 1):
                try:
                    img_array = np.array(image)
                    results = reader.readtext(img_array, detail=0)
                    page_text = "\n".join(results)
                    if page_text.strip():
                        all_text.append(f"[Page {page_num}]\n{page_text}")
                    
                    if page_num % 10 == 0:
                        logger.debug(f"Processed {page_num} pages")
                        
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num}: {e}")
                    continue
            
            combined = "\n".join(all_text)
            logger.info(f"OCR extraction complete: {len(combined)} chars from {len(images)} pages")
            return combined
            
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {e}")
            return ""

    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from PDF with intelligent fallback.
        1. Try PyPDF2 (fast, works for text PDFs)
        2. If minimal text, try OCR (slower, works for scanned PDFs)
        """
        logger.debug(f"Extracting text from PDF: {file_path}")
        
        # Try PyPDF2 first (fast path)
        text, success = self._extract_with_pypdf2(file_path)
        
        if success:
            return text
        
        # Fallback to OCR for scanned PDFs
        logger.info(f"PyPDF2 failed, trying OCR for {file_path.name}")
        ocr_text = self._extract_with_ocr(file_path)
        
        if ocr_text:
            return ocr_text
        
        # Both methods failed
        logger.error(f"Failed to extract text from {file_path} with both PyPDF2 and OCR")
        return ""


class DocxProcessor(DocumentProcessor):
    """Processor for Microsoft Word documents."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".docx"}

    def extract_text(self, file_path: Path) -> str:
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return ""

        try:
            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return ""


class DocProcessor(DocumentProcessor):
    """Processor for Legacy Microsoft Word documents (.doc)."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".doc"}

    def extract_text(self, file_path: Path) -> str:
        # 1️⃣ Tenta antiword (mais comum)
        try:
            result = subprocess.run(
                ["antiword", str(file_path)],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True
            )
            return result.stdout.decode("utf-8", errors="ignore")
        except FileNotFoundError:
            logger.warning("antiword not installed, trying next fallback...")
        except Exception as e:
            logger.warning(f"antiword failed for {file_path}: {e}")

        # 2️⃣ Tenta catdoc
        try:
            result = subprocess.run(
                ["catdoc", str(file_path)],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True
            )
            return result.stdout.decode("utf-8", errors="ignore")
        except FileNotFoundError:
            logger.warning("catdoc not installed, trying next fallback...")
        except Exception as e:
            logger.warning(f"catdoc failed for {file_path}: {e}")

        # 3️⃣ Fallback: converter via LibreOffice → docx
        try:
            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", str(file_path.parent),
                    str(file_path)
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True
            )

            converted = file_path.with_suffix(".docx")

            if converted.exists():
                from docx import Document
                doc = Document(converted)
                return "\n".join(p.text for p in doc.paragraphs)

        except FileNotFoundError:
            logger.error("LibreOffice not installed. Cannot process .doc file.")
        except Exception as e:
            logger.error(f"LibreOffice conversion failed for {file_path}: {e}")

        logger.error(f"All extraction methods failed for {file_path}")
        return ""

class CSVProcessor(DocumentProcessor):
    """Processor for CSV files."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".csv"

    def extract_text(self, file_path: Path) -> str:
        try:
            import csv
        except ImportError:
            logger.error("csv module not available")
            return ""

        try:
            text = []
            with open(file_path, "r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    text.append(str(row))
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from CSV {file_path}: {e}")
            return ""


class JSONProcessor(DocumentProcessor):
    """Processor for JSON files."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".json"

    def extract_text(self, file_path: Path) -> str:
        try:
            import json
        except ImportError:
            logger.error("json module not available")
            return ""

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return json.dumps(data, indent=2)
        except Exception as e:
            logger.error(f"Error extracting text from JSON {file_path}: {e}")
            return ""


class ImageProcessor(DocumentProcessor):
    """Processor for image files (using OCR if available)."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp"}

    def extract_text(self, file_path: Path) -> str:
        try:
            from PIL import Image
        except ImportError:
            logger.warning("Pillow not installed. Returning image metadata only.")
            return f"[Image: {file_path.name}]"

        try:
            # For now, just return metadata. Full OCR requires pytesseract
            img = Image.open(file_path)
            return f"[Image: {file_path.name} - {img.size} - {img.format}]"
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            return f"[Image: {file_path.name}]"


class ProcessorFactory:
    """Factory for creating appropriate document processors."""

    _processors = [
        PDFProcessor(),
        DocxProcessor(),
        DocProcessor(),
        CSVProcessor(),
        JSONProcessor(),
        ImageProcessor(),
        TextProcessor(),  # Keep last as fallback
    ]

    @classmethod
    def get_processor(cls, file_path: Path) -> Optional[DocumentProcessor]:
        """Get appropriate processor for a file."""
        for processor in cls._processors:
            if processor.can_process(file_path):
                return processor
        return None

    @classmethod
    def can_process(cls, file_path: Path) -> bool:
        """Check if any processor can handle the file."""
        return cls.get_processor(file_path) is not None

    @classmethod
    def register_processor(cls, processor: DocumentProcessor, priority: int = 0) -> None:
        """Register a custom processor."""
        if priority == 0:
            cls._processors.append(processor)
        else:
            cls._processors.insert(priority, processor)
